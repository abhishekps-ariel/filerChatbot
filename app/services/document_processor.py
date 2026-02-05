from pypdf import PdfReader
from docx import Document
from typing import List, Tuple
import io
import json


class SimpleTextSplitter:
    """Simple text splitter without external dependencies."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def split_text(self, text: str) -> List[str]:
        """Split text into chunks with overlap."""
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + self.chunk_size
            chunk = text[start:end]
            
            # Try to break at sentence or word boundary
            if end < text_length:
                # Look for sentence end
                last_period = chunk.rfind('. ')
                if last_period > self.chunk_size // 2:
                    end = start + last_period + 2
                    chunk = text[start:end]
                else:
                    # Look for word boundary
                    last_space = chunk.rfind(' ')
                    if last_space > self.chunk_size // 2:
                        end = start + last_space
                        chunk = text[start:end]
            
            chunks.append(chunk.strip())
            start = end - self.chunk_overlap
        
        return [c for c in chunks if c]  # Remove empty chunks


class DocumentProcessor:
    """Process PDF documents and split them into chunks."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = SimpleTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)
            
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from Word (DOCX) file."""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    def extract_text_from_markdown(self, file_content: bytes) -> str:
        """Extract text from Markdown (.md) file."""
        try:
            text = file_content.decode('utf-8')
            return text.strip()
        except Exception as e:
            raise ValueError(f"Failed to extract text from Markdown: {str(e)}")
    
    def extract_text_from_json(self, file_content: bytes) -> str:
        """Extract text from JSON file (Q&A dataset)."""
        try:
            data = json.loads(file_content.decode('utf-8'))
            
            # If it's a list of Q&A entries
            if isinstance(data, list):
                text_parts = []
                for entry in data:
                    if isinstance(entry, dict):
                        # Extract question and answer
                        question = entry.get('question', '')
                        answer = entry.get('answer', '')
                        category = entry.get('category', '')
                        intent = entry.get('intent', '')
                        
                        # Format as readable text
                        if question and answer:
                            text_parts.append(f"Category: {category}")
                            text_parts.append(f"Intent: {intent}")
                            text_parts.append(f"Q: {question}")
                            text_parts.append(f"A: {answer}")
                            text_parts.append("---")
                
                return "\n".join(text_parts)
            else:
                # Fallback: convert entire JSON to string
                return json.dumps(data, indent=2)
        except Exception as e:
            raise ValueError(f"Failed to extract text from JSON: {str(e)}")
    
    def chunk_text(self, text: str) -> List[str]:
        """Split text into chunks."""
        if not text:
            raise ValueError("Text is empty")
        
        chunks = self.text_splitter.split_text(text)
        return chunks
    
    def process_document(self, file_content: bytes, filename: str) -> Tuple[List[str], dict]:
        """
        Process document file (PDF or DOCX): extract text and chunk it.
        
        Returns:
            Tuple of (chunks, metadata)
        """
        # Determine file type and extract text
        if filename.lower().endswith('.pdf'):
            text = self.extract_text_from_pdf(file_content)
            file_type = "PDF"
        elif filename.lower().endswith('.docx'):
            text = self.extract_text_from_docx(file_content)
            file_type = "DOCX"
        elif filename.lower().endswith('.md'):
            text = self.extract_text_from_markdown(file_content)
            file_type = "Markdown"
        elif filename.lower().endswith('.json'):
            text = self.extract_text_from_json(file_content)
            file_type = "JSON"
        else:
            raise ValueError(f"Unsupported file type. Only PDF, DOCX, MD, and JSON are supported.")
        
        # Create chunks
        chunks = self.chunk_text(text)
        
        # Create metadata
        metadata = {
            "filename": filename,
            "file_type": file_type,
            "total_chunks": len(chunks),
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "total_characters": len(text)
        }
        
        return chunks, metadata
